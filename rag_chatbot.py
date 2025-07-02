# rag_chatbot.py
import os
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument
from docx import Document as DocxDocument
import openai
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch
from unstructured.partition.pdf import partition_pdf
from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# --- Load Groq API LLM ---
def get_llm():
    if not Config.GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY not found in environment variables")

    openai.api_key = Config.GROQ_API_KEY
    openai.api_base = "https://api.groq.com/openai/v1"
    model_name = "llama3-8b-8192"

    def query_llm(prompt, max_retries=3):
        for attempt in range(max_retries):
            try:
                response = openai.ChatCompletion.create(
                    model=model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.5,
                    max_tokens=512
                )
                return response['choices'][0]['message']['content']
            except Exception as e:
                logger.error(f"LLM error (attempt {attempt + 1}): {str(e)}")
                if attempt == max_retries - 1:
                    return "I apologize, but I'm having trouble processing your request. Please try again."

    return query_llm


# --- Embedding Model ---
_embedding_model = None


def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        model_name = "BAAI/bge-small-en-v1.5"
        _embedding_model = HuggingFaceEmbeddings(model_name=model_name)
    return _embedding_model


# --- Reranker Model ---
_tokenizer = None
_reranker_model = None


def load_reranker():
    global _tokenizer, _reranker_model
    if _tokenizer is None:
        _tokenizer = AutoTokenizer.from_pretrained("cross-encoder/ms-marco-MiniLM-L-6-v2")
        _reranker_model = AutoModelForSequenceClassification.from_pretrained("cross-encoder/ms-marco-MiniLM-L-6-v2")
    return _tokenizer, _reranker_model


def rerank(query, docs):
    if not docs:
        return []

    try:
        tokenizer, reranker_model = load_reranker()
        pairs = [(query, doc.page_content) for doc in docs]
        inputs = tokenizer(pairs, padding=True, truncation=True, return_tensors="pt", max_length=512)

        with torch.no_grad():
            scores = reranker_model(**inputs).logits.squeeze()

        if len(docs) == 1:
            scores = scores.unsqueeze(0)

        sorted_indices = torch.argsort(scores, descending=True)
        reranked_docs = [docs[i] for i in sorted_indices]
        return reranked_docs
    except Exception as e:
        logger.error(f"Reranking error: {str(e)}")
        return docs  # Return original docs if reranking fails


# --- PDF Loader and Splitter ---
def smart_load_pdf(file_path):
    try:
        elements = partition_pdf(file_path)
        docs = []
        for el in elements:
            if hasattr(el, 'text') and el.text and hasattr(el, 'category') and el.category in ["Title",
                                                                                               "NarrativeText"]:
                docs.append(LangchainDocument(page_content=el.text))
        return docs
    except Exception as e:
        logger.error(f"PDF loading error: {str(e)}")
        # Fallback to basic PDF reading
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        docs = []
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                docs.append(LangchainDocument(page_content=text))
        return docs


# --- DOCX Loader ---
def smart_load_docx(file_path):
    try:
        doc = DocxDocument(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [LangchainDocument(page_content=full_text)]
    except Exception as e:
        logger.error(f"DOCX loading error: {str(e)}")
        return []


def split_documents_smartly(docs):
    if not docs:
        return []

    try:
        md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=[("#", "heading")])
        structured_text = "\n\n".join(doc.page_content for doc in docs)
        structured_docs = md_splitter.split_text(structured_text)

        if len(structured_docs) > 5:
            return structured_docs
        else:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=700,
                chunk_overlap=100,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            return splitter.split_documents(docs)
    except Exception as e:
        logger.error(f"Document splitting error: {str(e)}")
        # Simple fallback splitting
        splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=100)
        return splitter.split_documents(docs)


# --- Load and Split Smartly ---
def load_and_split(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_ext = file_path.lower().split('.')[-1]

    if file_ext == "pdf":
        docs = smart_load_pdf(file_path)
    elif file_ext == "docx":
        docs = smart_load_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    return split_documents_smartly(docs)


# --- Vectorstore Builder ---
def build_vector_db(chunks):
    if not chunks:
        raise ValueError("No content found in the document")

    embedding_model = get_embedding_model()
    return Chroma.from_documents(chunks, embedding_model)


# --- RAG Answer Retriever ---
def retrieve_answer(file_path, query):
    try:
        chunks = load_and_split(file_path)
        if not chunks:
            return "I couldn't extract any content from the uploaded file. Please check if the file contains readable text."

        vectordb = build_vector_db(chunks)
        retriever = vectordb.as_retriever(search_kwargs={"k": min(10, len(chunks))})

        retrieved_docs = retriever.get_relevant_documents(query)
        if not retrieved_docs:
            return "I couldn't find relevant information in the document to answer your question."

        reranked_docs = rerank(query, retrieved_docs)[:4]

        llm = get_llm()
        context = "\n\n".join(doc.page_content for doc in reranked_docs)

        prompt = f"""
You are a helpful assistant. Answer the user's question based on the context below.
Always respond in the same language the question was asked.
If the context doesn't contain enough information to answer the question, say so clearly.

Context:
{context[:3000]}  # Limit context length

Question:
{query}

Answer:
"""

        return llm(prompt)
    except Exception as e:
        logger.error(f"Answer retrieval error: {str(e)}")
        return f"I encountered an error while processing your document: {str(e)}"


# --- Summarize Full Document ---
def summarize_document(file_path):
    try:
        chunks = load_and_split(file_path)
        if not chunks:
            return "I couldn't extract any content from the uploaded file. Please check if the file contains readable text."

        llm = get_llm()
        # Limit text length to avoid token limits
        full_text = "\n\n".join([chunk.page_content for chunk in chunks])[:5000]

        prompt = f"""
You are a professional summarizer. 
Summarize the following document into **three key sections**:
1. Main Topics
2. Important Details
3. Conclusion

Use clear and simple language. 

Here is the document:
---
{full_text}
---
"""
        return llm(prompt)
    except Exception as e:
        logger.error(f"Summarization error: {str(e)}")
        return f"I encountered an error while summarizing your document: {str(e)}"


# --- File Cleaner ---
def remove_temp_file(file_path):
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Removed temporary file: {file_path}")
    except Exception as e:
        logger.error(f"Error removing file {file_path}: {str(e)}")